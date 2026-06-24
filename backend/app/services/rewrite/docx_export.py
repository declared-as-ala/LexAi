"""DOCX export for revised contract text — structured legal layout."""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.rewrite.export_clean import prepare_export_body

# ── Heading detection (mirrors pdf_export logic) ────────────────────

def _classify_line(line: str) -> str:
    s = line.strip()
    if not s:
        return "spacer"
    if re.match(r"^(ARTICLE|CHAPITRE|TITRE|SECTION|CLAUSE)\s+\d+", s, re.IGNORECASE):
        return "h1"
    if s.isupper() and 4 <= len(s) <= 90 and not s.startswith(("-", "•", "*")):
        return "h1"
    if re.match(r"^(\d+\.)+\d*\s+\S", s) or re.match(r"^[IVXLC]+\.\s+\S", s):
        return "h2"
    if re.match(r"^[a-zA-Z]\)\s+\S", s) or re.match(r"^[a-zA-Z]\.\s+\S", s):
        return "h3"
    if s.endswith(":") and len(s) <= 80:
        return "h3"
    if s.startswith(("- ", "– ", "• ", "* ", "· ")):
        return "list_item"
    return "body"


def _parse_blocks(text: str) -> list[tuple[str, str]]:
    blocks: list[tuple[str, str]] = []
    prev_spacer = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        kind = _classify_line(line)
        if kind == "spacer":
            if not prev_spacer and blocks:
                blocks.append(("spacer", ""))
            prev_spacer = True
            continue
        prev_spacer = False
        s = line.strip()
        if (
            kind == "body"
            and blocks
            and blocks[-1][0] == "body"
            and not blocks[-1][1].endswith((".", ":", ";", "?", "!"))
        ):
            blocks[-1] = ("body", blocks[-1][1] + " " + s)
        else:
            blocks.append((kind, s))
    return blocks


# ── Word field helpers ───────────────────────────────────────────────

def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    r = run._element
    for tag, ftype, text in [
        ("w:fldChar", "begin", None),
        ("w:instrText", None, " PAGE "),
        ("w:fldChar", "separate", None),
        ("w:t", None, "1"),
        ("w:fldChar", "end", None),
    ]:
        el = OxmlElement(tag)
        if ftype:
            el.set(qn("w:fldCharType"), ftype)
        if tag == "w:instrText":
            el.set(qn("xml:space"), "preserve")
        if text:
            el.text = text
        r.append(el)


def _shade_paragraph(paragraph, fill_hex: str = "E8F5E9") -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    paragraph._element.get_or_add_pPr().append(shd)


def _set_font(run, name: str = "Times New Roman", size: float = 12.0,
              bold: bool = False, color: tuple | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _changed_applied_texts(metadata: list[dict[str, Any]] | None) -> set[str]:
    if not metadata:
        return set()
    out: set[str] = set()
    for m in metadata:
        if not m.get("changed"):
            continue
        t = (m.get("applied_text") or "").strip()
        if len(t) >= 12:
            out.add(t)
    return out


# ── Document builder ─────────────────────────────────────────────────

def build_docx_bytes(
    display_title: str,
    body_text: str,
    revision_summary: list[str] | None = None,
    *,
    export_date: str | None = None,
    revision_metadata: list[dict[str, Any]] | None = None,
) -> bytes:
    when = export_date or datetime.now(timezone.utc).strftime("%Y-%m-%d")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.2)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.35
    pf.space_after = Pt(4)

    # ── Header ──
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ht = hp.add_run((display_title.strip() or "Contrat Révisé") + "  ·  " + when)
    _set_font(ht, size=9, bold=True, color=(120, 90, 10))

    # ── Footer ──
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Page ")
    _set_font(fr, size=9, color=(100, 100, 100))
    _add_page_number_field(fp)
    fr2 = fp.add_run("  —  LexAI · Contrat révisé")
    _set_font(fr2, size=9, color=(100, 100, 100))

    # ── Title block ──
    title_text = (display_title.strip() or "Contrat Révisé").upper()
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(10)
    tp.paragraph_format.space_after = Pt(6)
    tr = tp.add_run(title_text)
    _set_font(tr, size=15, bold=True, color=(18, 18, 26))

    # Date line
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dp.paragraph_format.space_after = Pt(16)
    dr = dp.add_run(f"Document exporté le {when} via LexAI")
    _set_font(dr, size=9, color=(120, 120, 120))

    # Separator paragraph (thin border effect via shading)
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(12)
    _shade_paragraph(sep, fill_hex="C9A336")  # gold line

    body = prepare_export_body(body_text)
    highlight_set = _changed_applied_texts(revision_metadata)
    blocks = _parse_blocks(body)

    for kind, text in blocks:
        if kind == "spacer":
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(2)
            continue

        is_changed = any(hl in text for hl in highlight_set) if highlight_set else False

        if kind == "h1":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(text)
            _set_font(r, size=12, bold=True, color=(18, 18, 26))
            # Thin gold underline via bottom border
            ppr = p._element.get_or_add_pPr()
            pb = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "2")
            bot.set(qn("w:color"), "C9A336")
            pb.append(bot)
            ppr.append(pb)

        elif kind == "h2":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(text)
            _set_font(r, size=11.5, bold=True, color=(30, 50, 100))

        elif kind == "h3":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            r = p.add_run(text)
            _set_font(r, size=11, bold=True, color=(50, 70, 130))

        elif kind == "list_item":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(2)
            # Remove the bullet prefix chars since Word adds its own
            clean = re.sub(r"^[-–•*·]\s+", "", text)
            r = p.add_run(clean)
            _set_font(r, size=11)
            if is_changed:
                _shade_paragraph(p)

        else:  # body
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.first_line_indent = Cm(0.5)
            r = p.add_run(text)
            _set_font(r, size=11)
            if is_changed:
                _shade_paragraph(p)

    # ── Revision annex ──
    if revision_summary:
        doc.add_page_break()
        ann = doc.add_paragraph()
        ann.paragraph_format.space_after = Pt(8)
        ar = ann.add_run("ANNEXE — TRACE DES MODIFICATIONS")
        _set_font(ar, size=12, bold=True, color=(18, 18, 26))
        for line in revision_summary:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(line)
            _set_font(r, size=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def revision_summary_lines(metadata: list[dict[str, Any]], max_items: int = 50) -> list[str]:
    lines: list[str] = []
    for entry in metadata[:max_items]:
        if not entry.get("changed"):
            continue
        cid = entry.get("clause_id") or "?"
        lines.append(f"Clause {cid} : texte remplacé (positions {entry.get('start_char')}–{entry.get('end_char')})")
    if not lines:
        lines.append("Aucune clause modifiée dans cette version.")
    return lines
